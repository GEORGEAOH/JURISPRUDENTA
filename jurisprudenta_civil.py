import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
import io
import datetime

# 1. Configurare Pagină
st.set_page_config(page_title="Expert Jurisprudență GeorgeAOH", page_icon="⚖️", layout="wide")

# 2. Funcție pentru Calcul Suprafață din Coordonate Stereo 70
def calculeaza_suprafata_stereo70(puncte):
    n = len(puncte)
    area = 0.0
    for i in range(n):
        j = (i + 1) % n
        area += puncte[i][0] * puncte[j][1]
        area -= puncte[j][0] * puncte[i][1]
    return abs(area) / 2.0

# 3. Stil Vizual
st.markdown("""<style>.stCheckbox { background-color: #f0f2f6; padding: 5px; border-radius: 5px; }</style>""", unsafe_allow_html=True)

# 4. Bara Laterală - Identificare Dosar și Termen
with st.sidebar:
    st.header("📝 Detalii Dosar")
    nr_dosar = st.text_input("Dosar nr:", value="5975/221/2018")
    instanta = st.text_input("Instanța:", value="Tribunalul Hunedoara - Secția I Civilă")
    termen_jud = st.date_input("📅 Termen Judecată:", value=datetime.date(2026, 4, 14))
    
    st.divider()
    st.subheader("🏛️ Situație Intervenienți")
    st.info("Proprietarii Parcelelor Top 1 și Top 2 au intervenit în Apel, susținând respectarea Planului Parcelar.")
    
    st.divider()
    btn_generate = st.button("🚀 GENEREAZĂ MEMORIUL WORD", use_container_width=True)

# 5. Corp Principal
st.title("⚖️ Expert Jurisprudență GeorgeAOH")
st.subheader(f"Obiecțiuni Răspuns nr. 6 | Obiectiv 1: Repoziționare Plan Parcelar")

# --- MODUL 1: TABEL COMPARATIV CELE 6 PARCELE ---
st.markdown("### 📊 Tabel Comparativ: Titlu vs. Propunere Expert")
date_parcele = {
    "Nr. Parcela": ["Top 1", "Top 2", "Top 3 (932mp)", "Top 4", "Top 5", "Top 6/1"],
    "Suprafata Titlu (mp)": [0.0, 0.0, 932.0, 0.0, 0.0, 0.0],
    "Suprafata Expert (mp)": [0.0, 0.0, 0.0, 0.0, 0.0, 0.0],
    "Status": ["Intervenient", "Intervenient/Translatată", "Reclamant/Diminuată", "Reclamant", "Diminuată", "Diminuată"]
}
df_parcele = pd.DataFrame(date_parcele)
edited_df = st.data_editor(df_parcele, num_rows="fixed", use_container_width=True)

# --- MODUL 2: CALCULATOR COORDONATE PENTRU PROBĂ ---
with st.expander("📐 Calculator Matematic Coordonate (X, Y)"):
    num_pct = st.number_input("Număr puncte contur:", min_value=3, value=4)
    coords = []
    c1, c2 = st.columns(2)
    for i in range(num_pct):
        with c1: x = st.number_input(f"Punct {i+1} X:", format="%.3f", key=f"x{i}")
        with c2: y = st.number_input(f"Punct {i+1} Y:", format="%.3f", key=f"y{i}")
        coords.append((x, y))
    
    if st.button("Calculează Suprafața"):
        aria = calculeaza_suprafata_stereo70(coords)
        st.write(f"📏 Suprafața rezultată din coordonate: **{aria:.2f} mp**")

# --- MODUL 3: JURISPRUDENȚA IMBATABILĂ ---
st.divider()
st.markdown("### 🏛️ Argumente Juridice Cheie (Muniția de Dosar)")
col_a, col_b = st.columns(2)

with col_a:
    st.error("**DECIZIA RIL nr. 53/2007 ÎCCJ**")
    st.write("Stabilește că grănițuirea este declarativă. Nu poate anula **Revendicarea (S. 832/2000)**. Grănițuirea vecinului cu 19% în plus este nelegală.")

with col_b:
    st.error("**DECIZIA nr. 27/2022 ÎCCJ**")
    st.write("Planul Parcelar și Titlurile L.18 sunt **acte de autoritate**. Expertul NU le poate modifica sau diminua pe cale incidentală.")

# --- 6. LOGICA GENERARE WORD ---
if btn_generate:
    doc = Document()
    
    # Antet
    p_header = doc.add_paragraph()
    run_h = p_header.add_run(f"CĂTRE: {instanta}\nDOSAR NR: {nr_dosar}\nTERMEN: {termen_jud.strftime('%d.%m.%Y')}")
    run_h.bold = True
    p_header.alignment = 2

    doc.add_heading('OBIECTIUNI LA RĂSPUNSUL LA OBIECTIUNI NR. 6', 0).alignment = 1
    doc.add_paragraph("Privind OBIECTIVUL NR. 1: Repoziționarea conform Planului Parcelar Validat.").italic = True

    # 1. Prioritatea Revendicării
    doc.add_heading('1. ÎNTÂIETATEA REVENDICĂRII (DECIZIA RIL 53/2007)', level=1)
    p1 = doc.add_paragraph()
    p1.add_run("Invocăm DECIZIA RIL nr. 53/2007 a ÎCCJ. ").bold = True
    p1.add_run(
        "Grănițuirea pârâtului (S. 1500/2001) prin care acesta și-a însușit un excedent de 19% nu poate înfrânge "
        "Autoritatea de Lucru Judecat a Sentinței în Revendicare nr. 832/2000. Grănițuirea nu este mod de dobândire a proprietății."
    )

    # 2. Planul Parcelar
    doc.add_heading('2. INTANGIBILITATEA PLANULUI PARCELAR (DECIZIA 27/2022)', level=1)
    p2 = doc.add_paragraph()
    p2.add_run("Invocăm DECIZIILE 27/2022 și 66/2020 ale ÎCCJ. ").bold = True
    p2.add_run(
        "Expertul a procedat la o translatare nelegală a parcelelor Top 3 și 4 peste Top 2 și la diminuarea suprafețelor "
        "din Titlu pentru Top 1, 3, 5 și 6/1. Planul Parcelar validat în procedura Legii 18/1991 este obligatoriu și "
        "nu poate fi eludat prin măsurători tehnice care ignoră actele de autoritate."
    )

    # 3. Intervenția și Tabelul
    doc.add_heading('3. ANALIZA DISCREPANȚELOR ȘI VĂTĂMAREA TERȚILOR', level=1)
    doc.add_paragraph(
        "Manevra de translatare a expertului a generat intervenția proprietarilor Top 1 și Top 2 în Apel, "
        "demonstrând destabilizarea întregii tarlale. Tabelul de mai jos evidențiază diminuările nelegale:"
    )
    
    # Adăugare Tabel în Word
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nr. Top'
    hdr_cells[1].text = 'Suprafață Titlu (mp)'
    hdr_cells[2].text = 'Suprafață Expert (mp)'

    for _, row in edited_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Nr. Parcela'])
        row_cells[1].text = str(row['Suprafata Titlu (mp)'])
        row_cells[2].text = str(row['Suprafata Expert (mp)'])

    doc.add_paragraph("\nSOLICITĂM: Respingerea răspunsului expertului, refacerea raportului și respectarea suprafețelor din Titluri (Art. 187 C.pc.).")

    # Finalizare fișier
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    st.sidebar.download_button(
        label="📥 DESCARCĂ MEMORIUL FINAL",
        data=bio,
        file_name=f"Obiectiuni_GeorgeAOH_{nr_dosar.replace('/','_')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    st.success("✅ Documentul a fost generat! Verifică butonul din stânga.")
